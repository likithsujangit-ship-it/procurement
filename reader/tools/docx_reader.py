"""
Word Document (DOCX/DOC) Content Extraction Module.
Reads paragraphs, tables, and embedded images from Word files using python-docx and pytesseract.
"""

import io
from pathlib import Path
import docx
from tools.utils import setup_logger, preprocess_image_for_ocr

logger = setup_logger("docx_reader")


def extract_docx_text(filepath: Path) -> str:
    """
    Extracts text from paragraphs and tables in a .docx file, and runs OCR on embedded images.
    Provides troubleshooting fallback for older .doc files.
    
    Args:
        filepath: Path to the Word document on disk.
        
    Returns:
        Extracted text content as a string.
    """
    logger.info(f"Extracting text from DOCX: {filepath.name}")
    
    # Handle older .doc files gracefully
    if filepath.suffix.lower() == ".doc":
        logger.warning(f"File {filepath.name} is a legacy .doc file. python-docx only supports .docx.")
        return (
            "[Unsupported Format: Legacy .doc file]\n"
            "Troubleshooting: Please convert this file to the modern .docx format "
            "using Microsoft Word or LibreOffice to allow automated content reading."
        )
        
    native_chars = 0
    ocr_chars = 0
    
    try:
        doc = docx.Document(filepath)
        text_parts = []
        
        # 1. Extract paragraphs
        para_texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                para_texts.append(para.text.strip())
                native_chars += len(para.text.strip())
        if para_texts:
            text_parts.append("\n".join(para_texts))
                
        # 2. Extract tables if present
        table_parts = []
        for t_idx, table in enumerate(doc.tables):
            table_parts.append(f"--- Table {t_idx + 1} ---")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                row_str = " | ".join(row_data)
                if row_str.strip():
                    table_parts.append(row_str)
                    native_chars += len(row_str)
        if table_parts:
            text_parts.append("\n".join(table_parts))

        # 3. Extract and OCR embedded images
        try:
            import pytesseract
            from PIL import Image
            has_ocr = True
        except ImportError:
            has_ocr = False

        if has_ocr and hasattr(doc, "part") and hasattr(doc.part, "rels"):
            img_idx = 1
            for rel in doc.part.rels.values():
                if hasattr(rel, "target_part") and rel.target_part:
                    content_type = getattr(rel.target_part, "content_type", "")
                    if "image" in content_type:
                        try:
                            image_bytes = rel.target_part.blob
                            pil_img = Image.open(io.BytesIO(image_bytes))
                            processed_img = preprocess_image_for_ocr(pil_img)
                            ocr_text = pytesseract.image_to_string(processed_img).strip()
                            if ocr_text:
                                ocr_chars += len(ocr_text)
                                text_parts.append(f"\n[OCR from embedded image {img_idx}]: {ocr_text}")
                                img_idx += 1
                        except Exception as img_err:
                            logger.debug(f"Failed to process DOCX embedded image {img_idx}: {img_err}")

        logger.debug(f"DOCX Extraction Audit for {filepath.name}: native_chars={native_chars}, ocr_chars={ocr_chars}")
                
        if not text_parts:
            return "[DOCX Document is empty]"
            
        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error(f"Error reading DOCX {filepath.name}: {e}")
        return f"[Error reading DOCX file: {e}]"
